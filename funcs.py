import os
from docx import Document
import pandas as pd
import json

def read_file(fn):
    basename, ext = os.path.splitext(fn)
    if ext == '.docx':
        f = open(fn, 'rb')
        document = Document(f)
        f.close()
        txt = ''
        for para in document.paragraphs:
            try:
                txt = txt + para.text
            except:
                pass
    elif ext == '.txt':
        with open(fn, 'r', errors='ignore', encoding="utf-8") as f:
            txt = f.read()
    elif ext == ".xlsx":
        DF = pd.read_excel(fn)
        txt = DF.to_json(orient="columns")
    else:
        print('Unknown file type: skipping')
        txt = None
    return txt

def write_file(fn, txt):
    basename, ext = os.path.splitext(fn)
    if ext == ".txt":
        with open(fn, 'w', encoding="utf-8") as f:
            f.write(txt)
    elif ext == '.docx':
        document = Document()
        p = document.add_paragraph(txt)
        document.save(fn)
    elif ext == ".xlsx": # Currently: Will have been converted to a json string
        J = json.loads(txt)
        DF = pd.DataFrame.from_dict(J, orient='columns')
        DF.reset_index(level=0, inplace=True)
        DF.to_excel(fn)
    else:
        print("Unknown file type: Cannot save")
